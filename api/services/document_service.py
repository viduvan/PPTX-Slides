"""
Document Service — Logic for reading and parsing Word (.docx) and PDF (.pdf) files.
Developed by ChimSe (viduvan) - https://github.com/viduvan
Extracted from odin_slides/utils.py for API usage.
"""
import logging
from pathlib import Path
from typing import Callable, Awaitable

from docx import Document

from ..core.config import settings

logger = logging.getLogger("pptx_api.docs")


def read_docx(file_path: str | Path) -> str:
    """
    Read the full text content of a Word document.

    Args:
        file_path: Path to the .docx file.

    Returns:
        Full text content as a single string.
    """
    try:
        doc = Document(str(file_path))
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        return full_text
    except Exception as e:
        logger.error(f"Error reading Word file: {e}")
        raise


def read_pdf(file_path: str | Path) -> str:
    """
    Read the full text content of a PDF document.

    Args:
        file_path: Path to the .pdf file.

    Returns:
        Full text content as a single string.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(file_path))
        full_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                full_text += text + "\n"
        return full_text
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        raise


def read_big_docx(file_path: str | Path, chunk_size: int) -> list[str]:
    """
    Read a large Word document in chunks.

    Args:
        file_path: Path to the .docx file.
        chunk_size: Maximum number of words per chunk.

    Returns:
        List of text chunks.
    """
    try:
        doc = Document(str(file_path))
        chunks = []
        current_text = ""

        for paragraph in doc.paragraphs:
            paragraph_text = paragraph.text + "\n"
            if len(current_text.split()) + len(paragraph_text.split()) > chunk_size:
                if current_text.strip():
                    chunks.append(current_text)
                current_text = paragraph_text
            else:
                current_text += paragraph_text

        if current_text.strip():
            chunks.append(current_text)

        return chunks
    except Exception as e:
        logger.error(f"Error reading big Word file: {e}")
        raise


def read_big_pdf(file_path: str | Path, chunk_size: int) -> list[str]:
    """
    Read a large PDF document in chunks.

    Args:
        file_path: Path to the .pdf file.
        chunk_size: Maximum number of words per chunk.

    Returns:
        List of text chunks.
    """
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(str(file_path))
        chunks = []
        current_text = ""

        for page in reader.pages:
            page_text = page.extract_text()
            if not page_text:
                continue
            page_text += "\n"
            if len(current_text.split()) + len(page_text.split()) > chunk_size:
                if current_text.strip():
                    chunks.append(current_text)
                current_text = page_text
            else:
                current_text += page_text

        if current_text.strip():
            chunks.append(current_text)

        return chunks
    except Exception as e:
        logger.error(f"Error reading big PDF file: {e}")
        raise


def _split_text_into_chunks(text: str, chunk_size: int) -> list[str]:
    """
    Split already-extracted text into word-count-based chunks.

    Splits on paragraph boundaries (double newlines) to preserve structure.
    """
    paragraphs = text.split("\n")
    chunks = []
    current_text = ""

    for para in paragraphs:
        line = para + "\n"
        if len(current_text.split()) + len(line.split()) > chunk_size:
            if current_text.strip():
                chunks.append(current_text)
            current_text = line
        else:
            current_text += line

    if current_text.strip():
        chunks.append(current_text)

    return chunks


async def process_document(
    file_path: str | Path,
    summarize_fn: Callable[[str], Awaitable[str]],
) -> tuple[str, bool]:
    """
    Process a document: read it, and apply light summarization only for very large files.

    - Under 40,000 words: no summarization, send full content to LLM
    - 40,000 - 50,000 words: no summarization (Gemini 2.5 Flash handles this fine)
    - Over 50,000 words: light summarization (target 30-40% reduction)

    Args:
        file_path: Path to the .docx or .pdf file.
        summarize_fn: Async function to summarize text chunks.

    Returns:
        Tuple of (processed_content, was_summarized).
    """
    file_ext = Path(file_path).suffix.lower()
    if file_ext == ".pdf":
        word_content = read_pdf(file_path)
    else:
        word_content = read_docx(file_path)

    word_count = len(word_content.split())
    logger.info(f"Document read: {word_count} words")

    # Under threshold: send full content, no summarization needed
    if word_count <= settings.MAX_WORD_COUNT_WITHOUT_SUMMARIZATION:
        logger.info(f"Document under {settings.MAX_WORD_COUNT_WITHOUT_SUMMARIZATION} words, sending full content")
        return word_content, False

    # ── Light summarization for very large documents ─────────
    logger.info(
        f"Document has {word_count} words (>{settings.MAX_WORD_COUNT_WITHOUT_SUMMARIZATION}), "
        f"applying light summarization..."
    )

    chunk_size = settings.SUMMARIZE_CHUNK_SIZE

    # Split into chunks
    if file_ext == ".pdf":
        chunks = read_big_pdf(file_path, chunk_size)
    else:
        chunks = read_big_docx(file_path, chunk_size)

    logger.info(f"Split into {len(chunks)} chunks of ~{chunk_size} words")

    # Summarize each chunk (light — one round only)
    summarized_chunks = []
    for i, chunk in enumerate(chunks, 1):
        chunk_words = len(chunk.split())
        logger.info(f"  Summarizing chunk {i}/{len(chunks)} ({chunk_words} words)")
        summary = await summarize_fn(chunk)
        summarized_chunks.append(summary)

    result = "\n\n".join(summarized_chunks)
    result_word_count = len(result.split())
    reduction_pct = ((word_count - result_word_count) / word_count * 100)
    logger.info(
        f"Summarization done: {word_count} → {result_word_count} words "
        f"({reduction_pct:.0f}% reduction)"
    )

    # Safety truncation (unlikely but just in case)
    if result_word_count > settings.MAX_CONTENT_FOR_LLM:
        logger.warning(
            f"Content still {result_word_count} words after summarization. "
            f"Truncating to {settings.MAX_CONTENT_FOR_LLM} words."
        )
        words = result.split()
        result = " ".join(words[:settings.MAX_CONTENT_FOR_LLM])

    return result, True

